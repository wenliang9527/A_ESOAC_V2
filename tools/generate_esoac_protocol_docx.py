# -*- coding: utf-8 -*-
"""
Generate ESOAC通讯协议.docx — aligned with protocol.h / 功能说明 §3.1,
with structured layout, typography, and ASCII schematic blocks.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ESOAC通讯协议.docx"

COLOR_PRIMARY = RGBColor(0x2F, 0x54, 0x96)
COLOR_ACCENT = RGBColor(0x44, 0x72, 0xC4)
HEADER_FILL = "D9E2F3"
DOC_VERSION = "2.1"
DOC_DATE = "2026-03-27"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_table_header_row(table, row_index: int = 0) -> None:
    row = table.rows[row_index]
    for cell in row.cells:
        set_cell_shading(cell, HEADER_FILL)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_PRIMARY


def set_run_font_east_asia(run, name_en: str, name_asia: str, size_pt: float) -> None:
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    r = run._element.rPr
    if r is None:
        return
    rfonts = r.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        r.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name_en)
    rfonts.set(qn("w:hAnsi"), name_en)
    rfonts.set(qn("w:eastAsia"), name_asia)


def apply_document_theme(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")

    for lvl in range(1, 4):
        try:
            h = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        h.font.color.rgb = COLOR_PRIMARY
        h.font.bold = True
        hr = h.element.get_or_add_rPr()
        hf = hr.find(qn("w:rFonts"))
        if hf is None:
            hf = OxmlElement("w:rFonts")
            hr.insert(0, hf)
        hf.set(qn("w:eastAsia"), "微软雅黑")
        hf.set(qn("w:ascii"), "Calibri Light")
        hf.set(qn("w:hAnsi"), "Calibri Light")
        if lvl == 1:
            h.font.size = Pt(16)
        elif lvl == 2:
            h.font.size = Pt(13)
        else:
            h.font.size = Pt(11)

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)


def add_spacer(doc: Document, lines: float = 0.25) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(lines * 12)


def add_callout(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    r0 = p.add_run(title + " ")
    set_run_font_east_asia(r0, "Calibri", "微软雅黑", 11)
    r0.bold = True
    r0.font.color.rgb = COLOR_ACCENT
    r1 = p.add_run(body)
    set_run_font_east_asia(r1, "Calibri", "微软雅黑", 11)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(8)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    r = run._element.rPr
    if r is not None:
        rf = r.rFonts
        if rf is None:
            rf = OxmlElement("w:rFonts")
            r.insert(0, rf)
        rf.set(qn("w:ascii"), "Consolas")
        rf.set(qn("w:hAnsi"), "Consolas")
        rf.set(qn("w:eastAsia"), "Microsoft YaHei UI")


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font_east_asia(run, "Calibri", "微软雅黑", 11)


def build_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, c in enumerate(row_data):
            cells[i].text = str(c)
    style_table_header_row(tbl, 0)
    add_spacer(doc, 0.2)


def add_cover_block(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ESOAC 通讯协议")
    set_run_font_east_asia(r, "Calibri Light", "微软雅黑", 28)
    r.bold = True
    r.font.color.rgb = COLOR_PRIMARY

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("技术规范  Technical Specification")
    set_run_font_east_asia(sr, "Calibri", "微软雅黑", 14)
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    for line in (
        f"文档版本：V{DOC_VERSION}",
        f"发布日期：{DOC_DATE}",
        "适用范围：ESOAC 空调节能终端（BLE / MQTT / UART 统一二进制帧）",
    ):
        mp = doc.add_paragraph(line)
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in mp.runs:
            set_run_font_east_asia(run, "Calibri", "微软雅黑", 11)

    doc.add_paragraph()
    box_title = doc.add_paragraph()
    box_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    btr = box_title.add_run("文档要点")
    set_run_font_east_asia(btr, "Calibri", "微软雅黑", 12)
    btr.bold = True
    btr.font.color.rgb = COLOR_PRIMARY

    add_bullet_list(
        doc,
        [
            "与固件 protocol.h、protocol.c 及《空调节能终端程序功能说明文档》§3.1 对齐。",
            "同一帧格式用于 BLE GATT、MQTT 载荷、UART；命令字 16 位，线路上为小端。",
            "校验为 8 位累加和（非 CRC）；业务数据最大 252 字节（MAX_DATA_LENGTH）。",
        ],
    )
    doc.add_page_break()


def main() -> None:
    doc = Document()
    apply_document_theme(doc)

    add_cover_block(doc)

    doc.add_heading("文档导读", level=1)
    doc.add_paragraph(
        "建议按下列顺序阅读：先掌握第 1 章帧格式与校验，再查阅第 2 章命令分域表，"
        "MQTT 与字节序分别见第 3、4 章；附录供与历史 Word 规范对照。"
    )
    add_bullet_list(
        doc,
        [
            "第 1 章  协议帧与校验和",
            "第 2 章  命令字与载荷（按功能分组）",
            "第 3 章  MQTT 配置 TLV",
            "第 4 章  响应状态码与多字节序",
            "第 5 章  传输通道与处理流程（示意图）",
            "附录 A  旧版命令码映射参考",
            "附录 B  修订记录",
        ],
    )
    add_spacer(doc)

    doc.add_heading("1  协议帧与校验和", level=1)
    add_callout(
        doc,
        "定义：",
        "data_length 表示从 Data Mark 起至 Data 区末尾的字节数，等于 3+N（N 为业务载荷长度）。"
        "字段为 uint8，须 3+N≤255，故 N 最大 252。",
    )

    doc.add_heading("1.1  线格式与偏移", level=2)
    add_code_block(
        doc,
        [
            "字节偏移   0      1      2           3          4-5           6 … (6+N-1)    最后 1 字节",
            "         +------+------+-----------+----------+-------------+--------------+-----------+",
            "         | 5A   | 7A   |data_length|data_mark | command(LE) |   Data (N)   | checksum  |",
            "         +------+------+-----------+----------+-------------+--------------+-----------+",
            "           帧头(固定)    L=3+N      请求/响应/上报   Cmd 小端      业务数据      8 位累加和",
        ],
    )
    doc.add_paragraph("整帧长度（字节）= 4 + data_length。")
    doc.add_paragraph(
        "checksum 计算范围：帧头两字节 + data_length + data_mark + command 低字节 + command 高字节 + data 全部字节；"
        "结果为 8 位无符号累加和，不是 CRC。"
    )

    doc.add_heading("1.2  数据标记 data_mark", level=2)
    build_table(
        doc,
        ["取值", "名称", "说明"],
        [
            ("0x00", "请求", "上位机 → 设备"),
            ("0x01", "响应", "设备 → 上位机（应答某命令）"),
            ("0x02", "主动上报", "设备定时/事件推送"),
            ("0xFF", "错误", "错误类标记（与响应状态码配合使用）"),
        ],
    )

    doc.add_heading("2  命令字与载荷", level=1)
    doc.add_paragraph(
        "命令字为 16 位无符号数；线路上以低字节在前（小端）。下表按功能域分组，便于实现与评审。"
    )

    groups: list[tuple[str, list[tuple]]] = [
        (
            "2.1  基础与心跳",
            [
                ("0x0001", "心跳", "双向", "主动上报：4B 大端递增计数器"),
                ("0x0002", "设备信息", "查询", "应答：6B 设备ID + 4B 软件版本 + 4B 硬件版本"),
            ],
        ),
        (
            "2.2  空调控制",
            [
                ("0x0101", "设置开关", "设置", "data[0]: 0关 1开"),
                ("0x0102", "设置模式", "设置", "data[0]: 0自动 1制冷 2制热 3除湿 4送风 5睡眠"),
                ("0x0103", "设置温度", "设置", "data[0]: 16~30 ℃"),
                ("0x0104", "设置风速", "设置", "data[0]: 风速等级（枚举）"),
            ],
        ),
        (
            "2.3  状态与传感器",
            [
                ("0x0201", "获取状态", "查询", "应答 5B：开关/模式/温度/风速/连接"),
                ("0x0202", "获取功率", "查询", "应答 4B float 小端 IEEE754"),
                ("0x0203", "获取温度", "查询", "应答 4B float 小端 IEEE754"),
                ("0x0204", "时间同步", "设置", "8B：年16位大端+月日时分秒周各1B"),
                ("0x0205", "获取ADC", "查询", "应答 4B：两路 uint16 小端原始值"),
            ],
        ),
        (
            "2.4  红外",
            [
                ("0x0301", "红外学习开始", "设置", "data[0]: 按键索引"),
                ("0x0302", "红外学习停止", "设置", "无载荷"),
                ("0x0303", "红外发送", "设置", "data[0]: 按键索引"),
                ("0x0304", "红外学习结果", "—", "预留；固件未单独处理"),
                ("0x0305", "读取红外数据", "查询", "data[0]: 按键索引"),
                ("0x0306", "保存红外按键", "设置", "无载荷"),
            ],
        ),
        (
            "2.5  网络与设备名",
            [
                ("0x0401", "设置定时器", "—", "预留，未实现"),
                ("0x0402", "获取定时器", "—", "预留，未实现"),
                ("0x0501", "设置BLE名称", "设置", "data[0]=长度, 后跟名称字节"),
                ("0x0502", "获取BLE名称", "查询", "应答同上格式"),
                ("0x0503", "设置MQTT配置", "设置", "TLV，见第 3 章"),
                ("0x0504", "获取MQTT配置", "查询", "应答 TLV"),
            ],
        ),
        (
            "2.6  OTA 与通用应答",
            [
                ("0x0601", "OTA开始", "—", "预留，未实现"),
                ("0x0602", "OTA数据", "—", "预留，未实现"),
                ("0x0603", "OTA结束", "—", "预留，未实现"),
                ("0x8000", "CMD_RESPONSE", "响应", "data: [Cmd_H][Cmd_L][Status]，原命令字大端"),
            ],
        ),
    ]

    for title, cmd_rows in groups:
        doc.add_heading(title, level=2)
        build_table(doc, ["命令码", "名称", "方向", "数据格式 / 说明"], cmd_rows)

    doc.add_heading("3  MQTT 配置 TLV（0x0503 / 0x0504）", level=1)
    doc.add_paragraph("载荷由若干 TLV 顺序拼接组成，可只携带需要修改的字段（部分更新）。")
    add_code_block(
        doc,
        [
            "  [ type(1B) ][ len(1B) ][ value × len ] [ type ][ len ][ value ] …",
            "  示例：仅改服务器与端口 → 两个 TLV 即可，其余字段保持设备内原值。",
        ],
    )
    build_table(
        doc,
        ["type", "字段", "最大 len", "校验说明"],
        [
            ("0", "server_addr", "127", "可打印 ASCII，非空"),
            ("1", "server_port", "5", "可打印 ASCII，非空，如 1883"),
            ("2", "client_id", "127", "可打印 ASCII，非空"),
            ("3", "username", "127", "可打印 ASCII，允许空"),
            ("4", "password", "127", "允许特殊字符"),
            ("5", "subscribe_topic", "127", "可打印 ASCII，非空"),
            ("6", "publish_topic", "127", "可打印 ASCII，非空"),
        ],
    )
    add_callout(doc, "注意：", "蓝牙名称不属于 MQTT TLV，请使用命令 0x0501 / 0x0502。")

    doc.add_heading("4  响应状态码与多字节序", level=1)
    doc.add_heading("4.1  通用应答中的状态字节", level=2)
    build_table(
        doc,
        ["值", "符号（代码）", "含义"],
        [
            ("0x00", "STATUS_SUCCESS", "成功"),
            ("0x01", "STATUS_ERROR_PARAM", "参数错误"),
            ("0x02", "STATUS_ERROR_CMD", "未知命令"),
            ("0x03", "STATUS_ERROR_CHECKSUM", "帧累加和错误（非 CRC）"),
            ("0x04", "STATUS_ERROR_BUSY", "设备忙"),
            ("0x05", "STATUS_ERROR_STORAGE", "存储/Flash 错误"),
            ("0xFF", "STATUS_ERROR_FAIL", "执行失败"),
        ],
    )

    doc.add_heading("4.2  多字节字段字节序", level=2)
    build_table(
        doc,
        ["场景", "字节序", "备注"],
        [
            ("帧内 command（偏移 4–5）", "小端", "与 ARM 自然序一致"),
            ("CMD_RESPONSE 内嵌原命令字", "大端", "data[0]=高字节"),
            ("心跳 NOTIFY 计数器（4B）", "大端", "与帧内 command 约定不同"),
            ("CMD_SYNC_TIME 年份 16 位", "大端", "data[0] 为高字节"),
            ("float、ADC uint16", "小端", "IEEE754 / 原始采样值"),
        ],
    )

    doc.add_heading("5  传输通道与处理流程", level=1)
    doc.add_paragraph("同一解析入口保证 BLE、MQTT、UART 行为一致，便于测试与维护。")
    add_code_block(
        doc,
        [
            "  +----------+     +----------+     +----------------------+",
            "  | BLE GATT |     | MQTT     |     | UART1                |",
            "  +----+-----+     +----+-----+     +----------+-----------+",
            "       |                |                      |",
            "       +----------------+----------------------+",
            "                        |  原始字节流（同一帧格式）",
            "                        v",
            "               protocol_parse_frame()",
            "                        |",
            "                        v",
            "               protocol_verify_checksum",
            "                        |",
            "                        v",
            "               protocol_process_frame()  --> 命令分发 / 应答",
        ],
    )
    doc.add_paragraph(
        "固件源码路径：ESOACV2/ble_simple_peripheral/usercode/protocol.c、protocol.h。"
    )

    doc.add_page_break()
    doc.add_heading("附录 A  旧版 Word 命令码与固件对照", level=1)
    doc.add_paragraph(
        "若需兼容早期《ESOAC通讯协议》Word 中的编号，应在上位机或网关做映射，不建议静默修改设备端命令字。"
    )
    build_table(
        doc,
        ["旧版功能（摘要）", "旧版命令/类型", "当前固件"],
        [
            ("时间校准", "0x0053", "0x0204"),
            ("按键学习开始", "0x0010", "0x0301"),
            ("读取功率", "0x0005", "0x0202"),
            ("红外读键等", "0x0011 等", "0x0305 等"),
            ("学习键值保存", "0x0013", "0x0306"),
            ("读取键值", "0x0014", "0x0305"),
            ("空调开关", "0x0030", "0x0101"),
            ("温度/模式/风速", "0x0031~0x0034", "0x0103/0x0102/0x0104"),
            ("MQTT 服务器", "type 0x20", "TLV type 0"),
            ("端口 / 客户端 / 用户…", "0x21~0x24", "type 1~4"),
            ("订阅/发布", "0x25（重复）", "type 5 / 6"),
        ],
    )

    doc.add_heading("附录 B  修订记录", level=1)
    rev = doc.add_table(rows=1, cols=4)
    rev.style = "Table Grid"
    rh = rev.rows[0].cells
    for i, h in enumerate(("版本", "日期", "作者/角色", "说明")):
        rh[i].text = h
    style_table_header_row(rev, 0)
    rrow = rev.add_row().cells
    rrow[0].text = f"V{DOC_VERSION}"
    rrow[1].text = DOC_DATE
    rrow[2].text = "工程"
    rrow[3].text = "文档结构重组；表格样式与示意图；与 protocol.h 对齐"

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "本文档由 tools/generate_esoac_protocol_docx.py 生成；修改协议时请同步更新 protocol.h、"
        "《空调节能终端程序功能说明文档》§3.1 后重新运行脚本。"
    )
    for run in foot.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
